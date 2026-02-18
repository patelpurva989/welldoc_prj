"""Document text extraction and AI analysis service.

Phase 3: AI document integration for FDA Review Workflow.
Supports PDF and DOCX extraction, then uses Claude to summarize
the content in the context of an FDA regulatory submission.
"""
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# Maximum characters to send to AI to avoid token limits
MAX_EXTRACTED_CHARS = 10_000


# ============================================================================
# TEXT EXTRACTION
# ============================================================================

def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pdfplumber (primary) with PyPDF2 fallback."""
    text = ""

    # Try pdfplumber first - better layout handling
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages_text = []
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(f"[Page {page_num + 1}]\n{page_text}")
            text = "\n\n".join(pages_text)
        if text.strip():
            logger.info("PDF extracted via pdfplumber: %d chars from %s", len(text), file_path)
            return text
    except ImportError:
        logger.warning("pdfplumber not available, falling back to PyPDF2")
    except Exception as exc:
        logger.warning("pdfplumber extraction failed (%s), falling back to PyPDF2", exc)

    # Fallback: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        pages_text = []
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                pages_text.append(f"[Page {page_num + 1}]\n{page_text}")
        text = "\n\n".join(pages_text)
        logger.info("PDF extracted via PyPDF2: %d chars from %s", len(text), file_path)
        return text
    except ImportError:
        logger.error("Neither pdfplumber nor PyPDF2 is installed")
        raise RuntimeError(
            "PDF extraction libraries not available. "
            "Install pdfplumber or PyPDF2 in requirements.txt."
        )
    except Exception as exc:
        logger.error("PyPDF2 extraction failed for %s: %s", file_path, exc)
        raise RuntimeError(f"Failed to extract text from PDF: {exc}") from exc


def _extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Also extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    table_texts.append(row_text)

        all_text = "\n".join(paragraphs)
        if table_texts:
            all_text += "\n\n[TABLES]\n" + "\n".join(table_texts)

        logger.info("DOCX extracted: %d chars from %s", len(all_text), file_path)
        return all_text
    except ImportError:
        logger.error("python-docx is not installed")
        raise RuntimeError(
            "DOCX extraction library not available. "
            "Install python-docx in requirements.txt."
        )
    except Exception as exc:
        logger.error("DOCX extraction failed for %s: %s", file_path, exc)
        raise RuntimeError(f"Failed to extract text from DOCX: {exc}") from exc


async def extract_text_from_document(file_path: str, mime_type: str) -> str:
    """
    Extract text from a PDF or DOCX document.

    Args:
        file_path: Absolute path to the uploaded file on disk.
        mime_type: MIME type string of the file.

    Returns:
        Extracted text, truncated to MAX_EXTRACTED_CHARS to avoid token limits.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the MIME type is not supported for text extraction.
        RuntimeError: If extraction fails.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    pdf_types = {
        "application/pdf",
    }
    docx_types = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }

    if mime_type in pdf_types:
        raw_text = _extract_text_from_pdf(file_path)
    elif mime_type in docx_types:
        raw_text = _extract_text_from_docx(file_path)
    else:
        # Unsupported format for text extraction (e.g., images, XLSX)
        raise ValueError(
            f"Text extraction not supported for MIME type: {mime_type}. "
            "Supported formats: PDF, DOCX."
        )

    # Truncate to avoid exceeding model context / token limits
    if len(raw_text) > MAX_EXTRACTED_CHARS:
        logger.info(
            "Text truncated from %d to %d chars for %s",
            len(raw_text), MAX_EXTRACTED_CHARS, file_path
        )
        raw_text = raw_text[:MAX_EXTRACTED_CHARS] + "\n\n[... content truncated for AI analysis ...]"

    return raw_text


# ============================================================================
# AI SUMMARIZATION
# ============================================================================

async def analyze_document_with_ai(
    text: str,
    document_type: str,
    api_key: Optional[str] = None
) -> str:
    """
    Use Claude to generate a regulatory-focused summary of the extracted document text.

    Args:
        text: Extracted text from the document (should be pre-truncated).
        document_type: Category of the document (e.g., test_report, biocompatibility).
        api_key: Anthropic API key. Falls back to ANTHROPIC_API_KEY env var.

    Returns:
        AI-generated summary string focused on FDA regulatory relevance.
    """
    # Resolve API key
    resolved_key = api_key or os.environ.get("ANTHROPIC_API_KEY", "")
    if not resolved_key:
        logger.warning("ANTHROPIC_API_KEY not set - returning structured placeholder summary")
        return _generate_placeholder_summary(document_type, len(text))

    system_prompt = (
        "You are an FDA regulatory affairs expert. "
        "Your task is to read extracted text from a supporting document and produce "
        "a concise, structured summary (300-500 words) that highlights:\n"
        "1. Key findings and data points\n"
        "2. Regulatory relevance for a 510(k) or PMA submission\n"
        "3. Any compliance gaps or items needing attention\n"
        "4. How this document supports substantial equivalence claims\n\n"
        "Be precise, factual, and use regulatory language. "
        "Do not make up data not present in the text."
    )

    user_prompt = (
        f"Document Type: {document_type}\n\n"
        f"Extracted Content:\n{text}\n\n"
        "Please provide a regulatory-focused summary of this document."
    )

    try:
        from anthropic import Anthropic
        client = Anthropic(api_key=resolved_key)

        response = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=1024,
            messages=[
                {"role": "user", "content": user_prompt}
            ],
            system=system_prompt
        )

        summary = response.content[0].text
        logger.info(
            "AI summary generated for document_type=%s, input=%d chars, output=%d chars",
            document_type, len(text), len(summary)
        )
        return summary

    except ImportError:
        logger.error("anthropic package not installed")
        return _generate_placeholder_summary(document_type, len(text))
    except Exception as exc:
        logger.error("Anthropic API call failed: %s", exc)
        # Return placeholder rather than crashing the upload workflow
        return (
            f"[AI summary unavailable due to API error: {exc}]\n\n"
            + _generate_placeholder_summary(document_type, len(text))
        )


def _generate_placeholder_summary(document_type: str, text_length: int) -> str:
    """
    Return a structured placeholder summary when AI is unavailable.

    This gives downstream generation a valid non-empty string to include
    in the FDA submission prompt, making the placeholder easy to identify.
    """
    return (
        f"DOCUMENT SUMMARY (Placeholder - API key not configured)\n"
        f"Document Type: {document_type}\n"
        f"Extracted Text Length: {text_length} characters\n\n"
        f"Key Findings:\n"
        f"  - Document successfully uploaded and text extracted ({text_length} chars).\n"
        f"  - Full AI analysis requires a valid ANTHROPIC_API_KEY in the .env file.\n\n"
        f"Regulatory Relevance:\n"
        f"  - This {document_type} document should be reviewed manually until AI "
        f"summarization is enabled.\n\n"
        f"Next Steps:\n"
        f"  1. Set ANTHROPIC_API_KEY in docker-compose.yml or .env\n"
        f"  2. Re-trigger AI review via POST /api/v1/documents/{{id}}/ai-review\n"
        f"  3. The full summary will then be incorporated into FDA submission generation."
    )
